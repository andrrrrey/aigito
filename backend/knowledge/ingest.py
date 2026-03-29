"""
Document ingestion pipeline:
  file bytes → text extraction → chunks → OpenAI embeddings → Qdrant upsert
"""
import csv
import io
import logging
import os
import uuid
from typing import List

logger = logging.getLogger(__name__)

CHUNK_SIZE = 400     # words per chunk
CHUNK_OVERLAP = 40  # word overlap between chunks
EMBEDDING_MODEL = "text-embedding-3-small"
VECTOR_SIZE = 1536


# ── Text extraction ───────────────────────────────────────────────────────────

def extract_text(filename: str, content: bytes) -> str:
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else "txt"
    try:
        if ext == "txt":
            return content.decode("utf-8", errors="ignore")
        elif ext == "pdf":
            return _extract_pdf(content)
        elif ext == "docx":
            return _extract_docx(content)
        elif ext == "csv":
            return _extract_csv(content)
        else:
            return content.decode("utf-8", errors="ignore")
    except Exception as e:
        logger.warning(f"Text extraction failed for {filename}: {e}")
        return content.decode("utf-8", errors="ignore")


def _extract_pdf(content: bytes) -> str:
    import pypdf
    reader = pypdf.PdfReader(io.BytesIO(content))
    parts = []
    for page in reader.pages:
        text = page.extract_text()
        if text:
            parts.append(text)
    return "\n\n".join(parts)


def _extract_docx(content: bytes) -> str:
    import docx
    doc = docx.Document(io.BytesIO(content))
    parts = [para.text for para in doc.paragraphs if para.text.strip()]
    # Also extract tables
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                parts.append(row_text)
    return "\n\n".join(parts)


def _extract_csv(content: bytes) -> str:
    text = content.decode("utf-8", errors="ignore")
    reader = csv.reader(io.StringIO(text))
    rows = []
    headers = None
    for i, row in enumerate(reader):
        if i == 0:
            headers = row
            rows.append(" | ".join(row))
        else:
            if headers:
                rows.append(" | ".join(f"{h}: {v}" for h, v in zip(headers, row)))
            else:
                rows.append(" | ".join(row))
    return "\n".join(rows)


# ── Chunking ──────────────────────────────────────────────────────────────────

def chunk_text(text: str, chunk_size: int = CHUNK_SIZE, overlap: int = CHUNK_OVERLAP) -> List[str]:
    words = text.split()
    if not words:
        return []
    chunks = []
    i = 0
    while i < len(words):
        chunk_words = words[i : i + chunk_size]
        chunks.append(" ".join(chunk_words))
        i += chunk_size - overlap
    return [c for c in chunks if len(c.strip()) > 20]


# ── Qdrant helpers ────────────────────────────────────────────────────────────

def _get_qdrant():
    from qdrant_client import QdrantClient
    host = os.getenv("QDRANT_HOST", "qdrant")
    port = int(os.getenv("QDRANT_PORT", "6333"))
    return QdrantClient(host=host, port=port)


def _ensure_collection(client, collection_name: str):
    from qdrant_client.models import Distance, VectorParams
    existing = {c.name for c in client.get_collections().collections}
    if collection_name not in existing:
        client.create_collection(
            collection_name=collection_name,
            vectors_config=VectorParams(size=VECTOR_SIZE, distance=Distance.COSINE),
        )
        logger.info(f"Created Qdrant collection: {collection_name}")


# ── Main ingestion ────────────────────────────────────────────────────────────

async def ingest_document(company_id: str, doc_id: str, text: str) -> int:
    """
    Embed document chunks and upsert into Qdrant.
    Returns the number of chunks ingested.
    """
    chunks = chunk_text(text)
    if not chunks:
        return 0

    try:
        from openai import AsyncOpenAI
        from qdrant_client.models import PointStruct

        openai_api_key = os.getenv("OPENAI_API_KEY", "")
        openai_client = AsyncOpenAI(api_key=openai_api_key)
        qdrant = _get_qdrant()
        collection_name = f"company_{company_id}"
        _ensure_collection(qdrant, collection_name)

        # Embed in batches of 20
        batch_size = 20
        all_points = []
        for i in range(0, len(chunks), batch_size):
            batch = chunks[i : i + batch_size]
            response = await openai_client.embeddings.create(
                model=EMBEDDING_MODEL,
                input=batch,
            )
            for j, emb_data in enumerate(response.data):
                point = PointStruct(
                    id=str(uuid.uuid4()),
                    vector=emb_data.embedding,
                    payload={
                        "text": batch[j],
                        "doc_id": doc_id,
                        "company_id": company_id,
                        "chunk_index": i + j,
                    },
                )
                all_points.append(point)

        qdrant.upsert(collection_name=collection_name, points=all_points)
        logger.info(f"Ingested {len(all_points)} chunks for company {company_id}")
        return len(all_points)

    except Exception as e:
        logger.error(f"Ingestion failed for company {company_id}: {e}")
        return 0


async def delete_document_chunks(company_id: str, doc_id: str):
    """Remove all Qdrant points for a specific document."""
    try:
        from qdrant_client.models import Filter, FieldCondition, MatchValue
        qdrant = _get_qdrant()
        collection_name = f"company_{company_id}"
        existing = {c.name for c in qdrant.get_collections().collections}
        if collection_name not in existing:
            return
        qdrant.delete(
            collection_name=collection_name,
            points_selector=Filter(
                must=[FieldCondition(key="doc_id", match=MatchValue(value=doc_id))]
            ),
        )
        logger.info(f"Deleted chunks for doc {doc_id}")
    except Exception as e:
        logger.warning(f"Failed to delete doc chunks: {e}")


async def rebuild_company_index(company_id: str, documents: list[dict]) -> int:
    """Re-embed all documents for a company (drop collection first)."""
    try:
        qdrant = _get_qdrant()
        collection_name = f"company_{company_id}"
        existing = {c.name for c in qdrant.get_collections().collections}
        if collection_name in existing:
            qdrant.delete_collection(collection_name)
            logger.info(f"Dropped collection {collection_name} for rebuild")

        total = 0
        for doc in documents:
            n = await ingest_document(company_id, doc["id"], doc["content_text"] or "")
            total += n
        return total
    except Exception as e:
        logger.error(f"Rebuild failed for company {company_id}: {e}")
        return 0
